from datetime import datetime, timezone
import io
import requests
from html import escape
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

from docx import Document

from fastapi import HTTPException
from sqlalchemy.orm import Session
from google import genai

from app.core.config import settings
from app.models.audio import Audio
from app.models.proyecto import Proyecto
from app.models.transcripcion import Transcripcion
from app.models.solicitud import Solicitud
from app.models.solicitud_audio import SolicitudAudio
from app.models.resumen import Resumen


MODELO_RESUMEN = "gemini-2.5-flash"


def construir_prompt_resumen(texto: str, tipo_resumen: str) -> str:
    instrucciones = {
        "CORTO": "Genera un resumen breve, directo y claro en máximo 1 o 2 párrafos.",
        "MEDIO": "Genera un resumen intermedio, organizado en secciones cortas.",
        "DETALLADO": "Genera un resumen detallado, con temas principales, subtemas, ideas importantes y conclusiones."
    }

    return f"""
Eres un asistente académico especializado en resumir clases, reuniones, entrevistas y conferencias.

Tipo de resumen solicitado: {tipo_resumen}

Instrucciones:
{instrucciones.get(tipo_resumen, instrucciones["MEDIO"])}

Reglas:
- No inventes información.
- Usa solo el contenido de la transcripción.
- Mantén un lenguaje claro y ordenado.
- Si el texto no tiene suficiente información, indícalo.

Transcripción:
{texto}
"""


def obtener_url_ollama_activa() -> str:
    for host in ["ollama", "localhost", "127.0.0.1"]:
        try:
            url = f"http://{host}:11434/api/tags"
            response = requests.get(url, timeout=2)
            if response.status_code == 200:
                return f"http://{host}:11434"
        except Exception:
            continue
    return "http://ollama:11434"


def seleccionar_modelo_ollama(active_url: str) -> str:
    try:
        url = f"{active_url}/api/tags"
        response = requests.get(url, timeout=3)
        if response.status_code == 200:
            models = response.json().get("models", [])
            names = [m["name"] for m in models]
            generativos = [n for n in names if "embed" not in n]
            if generativos:
                for fav in ["llama3.2", "llama3", "mistral", "gemma", "phi3"]:
                    for g in generativos:
                        if fav in g:
                            return g
                return generativos[0]
    except Exception:
        pass
    return "llama3.2"


def generar_resumen_con_ollama(texto: str, tipo_resumen: str) -> str:
    active_url = obtener_url_ollama_activa()
    model = seleccionar_modelo_ollama(active_url)
    prompt = construir_prompt_resumen(texto, tipo_resumen)
    
    url = f"{active_url}/api/generate"
    payload = {
        "model": model,
        "prompt": prompt,
        "stream": False
    }
    response = requests.post(url, json=payload, timeout=300)
    response.raise_for_status()
    resumen = response.json()["response"].strip()
    if not resumen:
        raise ValueError("Ollama retornó un resumen vacío")
    return resumen


def generar_resumen_con_gemini(texto: str, tipo_resumen: str) -> str:
    client = genai.Client(api_key=settings.gemini_api_key)

    prompt = construir_prompt_resumen(texto, tipo_resumen)

    response = client.models.generate_content(
        model=MODELO_RESUMEN,
        contents=prompt
    )

    resumen = response.text

    if not resumen or not resumen.strip():
        raise HTTPException(
            status_code=500,
            detail="No se pudo generar el resumen"
        )

    return resumen.strip()


def generar_resumen_ia(texto: str, tipo_resumen: str) -> tuple[str, str]:
    """
    Genera resumen intentando usar Ollama y haciendo fallback a Gemini.
    Retorna la tupla (resumen_generado, modelo_usado).
    """
    try:
        texto_resumen = generar_resumen_con_ollama(texto, tipo_resumen)
        active_url = obtener_url_ollama_activa()
        model_name = seleccionar_modelo_ollama(active_url)
        return texto_resumen, f"Ollama ({model_name})"
    except Exception as e:
        print(f"Error con Ollama (se intentará fallback a Gemini): {e}")
        try:
            texto_resumen = generar_resumen_con_gemini(texto, tipo_resumen)
            return texto_resumen, f"Gemini ({MODELO_RESUMEN})"
        except Exception as gemini_err:
            raise HTTPException(
                status_code=500,
                detail=f"Error en IA de Ollama y Gemini (Fallback falló). Detalles: {str(gemini_err)}"
            )


def obtener_audio_usuario_con_transcripcion(
    db: Session,
    audio_id: int,
    usuario_id: int
):
    resultado = (
        db.query(Audio, Transcripcion)
        .join(Proyecto, Audio.id_proyecto == Proyecto.id)
        .join(Transcripcion, Transcripcion.id_audio == Audio.id)
        .filter(
            Audio.id == audio_id,
            Proyecto.id_usuario == usuario_id
        )
        .first()
    )

    if not resultado:
        raise HTTPException(
            status_code=404,
            detail="Audio no encontrado o todavía no tiene transcripción"
        )

    audio, transcripcion = resultado

    if not transcripcion.texto_generado:
        raise HTTPException(
            status_code=400,
            detail="La transcripción está vacía"
        )

    return audio, transcripcion


def obtener_transcripciones_proyecto(
    db: Session,
    proyecto_id: int,
    usuario_id: int
):
    proyecto = (
        db.query(Proyecto)
        .filter(
            Proyecto.id == proyecto_id,
            Proyecto.id_usuario == usuario_id
        )
        .first()
    )

    if not proyecto:
        raise HTTPException(
            status_code=404,
            detail="Proyecto no encontrado"
        )

    resultados = (
        db.query(Audio, Transcripcion)
        .join(Transcripcion, Transcripcion.id_audio == Audio.id)
        .filter(Audio.id_proyecto == proyecto_id)
        .order_by(Audio.creado_en.asc())
        .all()
    )

    if not resultados:
        raise HTTPException(
            status_code=400,
            detail="El proyecto no tiene audios transcritos"
        )

    return resultados


def crear_solicitud_resumen(
    db: Session,
    usuario_id: int,
    prompt_usado: str
) -> Solicitud:
    solicitud = Solicitud(
        id_usuario=usuario_id,
        tipo="RESUMEN",
        prompt_usado=prompt_usado,
        estado="PROCESANDO"
    )

    db.add(solicitud)
    db.commit()
    db.refresh(solicitud)

    return solicitud


def vincular_audios_solicitud(
    db: Session,
    solicitud_id: int,
    audio_ids: list[int]
):
    for audio_id in audio_ids:
        relacion = SolicitudAudio(
            id_solicitud=solicitud_id,
            id_audio=audio_id
        )
        db.add(relacion)

    db.commit()


def guardar_resumen(
    db: Session,
    solicitud: Solicitud,
    titulo: str,
    tipo_resumen: str,
    texto_resumen: str,
    audios_usados: list[int],
    modelo_usado: str = MODELO_RESUMEN
) -> Resumen:
    resumen = Resumen(
        id_solicitud=solicitud.id,
        titulo=titulo,
        tipo_resumen=tipo_resumen,
        contenido={
            "texto": texto_resumen,
            "audios_usados": audios_usados,
            "modelo_usado": modelo_usado
        }
    )

    db.add(resumen)

    solicitud.estado = "COMPLETADO"
    solicitud.mensaje_error = None
    solicitud.completado_en = datetime.now(timezone.utc)

    db.commit()
    db.refresh(resumen)

    return resumen


def generar_resumen_de_audio(
    db: Session,
    audio_id: int,
    usuario_id: int,
    tipo_resumen: str
) -> Resumen:
    audio, transcripcion = obtener_audio_usuario_con_transcripcion(
        db,
        audio_id,
        usuario_id
    )

    prompt = construir_prompt_resumen(
        transcripcion.texto_generado,
        tipo_resumen
    )

    solicitud = crear_solicitud_resumen(db, usuario_id, prompt)

    try:
        vincular_audios_solicitud(db, solicitud.id, [audio.id])

        texto_resumen, modelo_usado = generar_resumen_ia(
            transcripcion.texto_generado,
            tipo_resumen
        )

        return guardar_resumen(
            db=db,
            solicitud=solicitud,
            titulo=f"Resumen de {audio.titulo}",
            tipo_resumen=tipo_resumen,
            texto_resumen=texto_resumen,
            audios_usados=[audio.id],
            modelo_usado=modelo_usado
        )

    except Exception as e:
        solicitud.estado = "ERROR"
        solicitud.mensaje_error = str(e)
        db.commit()

        raise HTTPException(
            status_code=500,
            detail=f"Error al generar resumen: {str(e)}"
        )


def generar_resumen_de_proyecto(
    db: Session,
    proyecto_id: int,
    usuario_id: int,
    tipo_resumen: str
) -> Resumen:
    resultados = obtener_transcripciones_proyecto(
        db,
        proyecto_id,
        usuario_id
    )

    textos = []
    audio_ids = []

    for audio, transcripcion in resultados:
        if transcripcion.texto_generado:
            textos.append(f"Audio: {audio.titulo}\n{transcripcion.texto_generado}")
            audio_ids.append(audio.id)

    if not textos:
        raise HTTPException(
            status_code=400,
            detail="No existen transcripciones válidas para resumir"
        )

    texto_unificado = "\n\n---\n\n".join(textos)

    prompt = construir_prompt_resumen(texto_unificado, tipo_resumen)

    solicitud = crear_solicitud_resumen(db, usuario_id, prompt)

    try:
        vincular_audios_solicitud(db, solicitud.id, audio_ids)

        texto_resumen, modelo_usado = generar_resumen_ia(
            texto_unificado,
            tipo_resumen
        )

        return guardar_resumen(
            db=db,
            solicitud=solicitud,
            titulo="Resumen general del proyecto",
            tipo_resumen=tipo_resumen,
            texto_resumen=texto_resumen,
            audios_usados=audio_ids,
            modelo_usado=modelo_usado
        )

    except Exception as e:
        solicitud.estado = "ERROR"
        solicitud.mensaje_error = str(e)
        db.commit()

        raise HTTPException(
            status_code=500,
            detail=f"Error al generar resumen del proyecto: {str(e)}"
        )


def generar_resumen_seleccion(
    db: Session,
    usuario_id: int,
    titulo: str,
    audio_ids: list[int],
    tipo_resumen: str
) -> Resumen:
    if not audio_ids:
        raise HTTPException(
            status_code=400,
            detail="Debe seleccionar al menos un audio para generar el resumen"
        )

    textos = []
    for audio_id in audio_ids:
        audio, transcripcion = obtener_audio_usuario_con_transcripcion(
            db,
            audio_id,
            usuario_id
        )
        if transcripcion.texto_generado:
            textos.append(f"Audio: {audio.titulo}\n{transcripcion.texto_generado}")

    if not textos:
        raise HTTPException(
            status_code=400,
            detail="Ninguno de los audios seleccionados tiene transcripciones válidas"
        )

    texto_unificado = "\n\n---\n\n".join(textos)

    prompt = construir_prompt_resumen(texto_unificado, tipo_resumen)

    solicitud = crear_solicitud_resumen(db, usuario_id, prompt)

    try:
        vincular_audios_solicitud(db, solicitud.id, audio_ids)

        texto_resumen, modelo_usado = generar_resumen_ia(
            texto_unificado,
            tipo_resumen
        )

        resumen_titulo = titulo if titulo and titulo.strip() else f"Resumen de {len(audio_ids)} audios"

        return guardar_resumen(
            db=db,
            solicitud=solicitud,
            titulo=resumen_titulo,
            tipo_resumen=tipo_resumen,
            texto_resumen=texto_resumen,
            audios_usados=audio_ids,
            modelo_usado=modelo_usado
        )

    except Exception as e:
        solicitud.estado = "ERROR"
        solicitud.mensaje_error = str(e)
        db.commit()

        raise HTTPException(
            status_code=500,
            detail=f"Error al generar el resumen de selección: {str(e)}"
        )
    
def obtener_ultimo_resumen_audio(
    db: Session,
    audio_id: int,
    usuario_id: int
) -> Resumen:
    # Verificar que el audio pertenezca al usuario
    audio = (
        db.query(Audio)
        .join(Proyecto, Audio.id_proyecto == Proyecto.id)
        .filter(
            Audio.id == audio_id,
            Proyecto.id_usuario == usuario_id
        )
        .first()
    )

    if not audio:
        raise HTTPException(
            status_code=404,
            detail="Audio no encontrado"
        )

    resumen = (
        db.query(Resumen)
        .join(Solicitud, Resumen.id_solicitud == Solicitud.id)
        .join(SolicitudAudio, SolicitudAudio.id_solicitud == Solicitud.id)
        .filter(
            Solicitud.id_usuario == usuario_id,
            Solicitud.tipo == "RESUMEN",
            SolicitudAudio.id_audio == audio_id
        )
        .order_by(Resumen.creado_en.desc())
        .first()
    )

    if not resumen:
        raise HTTPException(
            status_code=404,
            detail="Este audio todavía no tiene resumen generado"
        )

    return resumen


def obtener_ultimo_resumen_proyecto(
    db: Session,
    proyecto_id: int,
    usuario_id: int
) -> Resumen:
    proyecto = (
        db.query(Proyecto)
        .filter(
            Proyecto.id == proyecto_id,
            Proyecto.id_usuario == usuario_id
        )
        .first()
    )

    if not proyecto:
        raise HTTPException(
            status_code=404,
            detail="Proyecto no encontrado"
        )

    resumen = (
        db.query(Resumen)
        .join(Solicitud, Resumen.id_solicitud == Solicitud.id)
        .join(SolicitudAudio, SolicitudAudio.id_solicitud == Solicitud.id)
        .join(Audio, SolicitudAudio.id_audio == Audio.id)
        .filter(
            Solicitud.id_usuario == usuario_id,
            Solicitud.tipo == "RESUMEN",
            Audio.id_proyecto == proyecto_id
        )
        .order_by(Resumen.creado_en.desc())
        .first()
    )

    if not resumen:
        raise HTTPException(
            status_code=404,
            detail="Este proyecto todavía no tiene resumen generado"
        )

    return resumen

def obtener_resumen_por_id_usuario(
    db: Session,
    resumen_id: int,
    usuario_id: int
) -> Resumen:
    resumen = (
        db.query(Resumen)
        .join(Solicitud, Resumen.id_solicitud == Solicitud.id)
        .filter(
            Resumen.id == resumen_id,
            Solicitud.id_usuario == usuario_id,
            Solicitud.tipo == "RESUMEN"
        )
        .first()
    )

    if not resumen:
        raise HTTPException(
            status_code=404,
            detail="Resumen no encontrado"
        )

    return resumen


def listar_resumenes_por_proyecto(
    db: Session,
    proyecto_id: int,
    usuario_id: int
):
    proyecto = (
        db.query(Proyecto)
        .filter(
            Proyecto.id == proyecto_id,
            Proyecto.id_usuario == usuario_id
        )
        .first()
    )

    if not proyecto:
        raise HTTPException(
            status_code=404,
            detail="Proyecto no encontrado"
        )

    resumenes = (
        db.query(Resumen)
        .join(Solicitud, Resumen.id_solicitud == Solicitud.id)
        .join(SolicitudAudio, SolicitudAudio.id_solicitud == Solicitud.id)
        .join(Audio, SolicitudAudio.id_audio == Audio.id)
        .filter(
            Solicitud.id_usuario == usuario_id,
            Solicitud.tipo == "RESUMEN",
            Audio.id_proyecto == proyecto_id
        )
        .order_by(Resumen.creado_en.desc())
        .distinct()
        .all()
    )

    return resumenes


def listar_resumenes_por_audio(
    db: Session,
    audio_id: int,
    usuario_id: int
):
    audio = (
        db.query(Audio)
        .join(Proyecto, Audio.id_proyecto == Proyecto.id)
        .filter(
            Audio.id == audio_id,
            Proyecto.id_usuario == usuario_id
        )
        .first()
    )

    if not audio:
        raise HTTPException(
            status_code=404,
            detail="Audio no encontrado"
        )

    resumenes = (
        db.query(Resumen)
        .join(Solicitud, Resumen.id_solicitud == Solicitud.id)
        .join(SolicitudAudio, SolicitudAudio.id_solicitud == Solicitud.id)
        .filter(
            Solicitud.id_usuario == usuario_id,
            Solicitud.tipo == "RESUMEN",
            SolicitudAudio.id_audio == audio_id
        )
        .order_by(Resumen.creado_en.desc())
        .all()
    )

    return resumenes


def obtener_texto_resumen(resumen: Resumen) -> str:
    if not resumen.contenido:
        return ""

    texto = resumen.contenido.get("texto", "")

    if not texto:
        return ""

    return texto


def generar_archivo_txt(resumen: Resumen) -> io.BytesIO:
    contenido = obtener_texto_resumen(resumen)

    texto = f"""
{resumen.titulo or "Resumen generado"}

Tipo de resumen: {resumen.tipo_resumen}
Fecha de creación: {resumen.creado_en}

Contenido:
{contenido}
"""

    archivo = io.BytesIO()
    archivo.write(texto.encode("utf-8"))
    archivo.seek(0)

    return archivo


def generar_archivo_pdf(resumen: Resumen) -> io.BytesIO:
    contenido = obtener_texto_resumen(resumen)

    archivo = io.BytesIO()

    doc = SimpleDocTemplate(
        archivo,
        pagesize=letter,
        rightMargin=50,
        leftMargin=50,
        topMargin=50,
        bottomMargin=50
    )

    styles = getSampleStyleSheet()
    elementos = []

    titulo = resumen.titulo or "Resumen generado"

    elementos.append(Paragraph(escape(titulo), styles["Title"]))
    elementos.append(Spacer(1, 12))

    elementos.append(Paragraph(f"<b>Tipo de resumen:</b> {escape(resumen.tipo_resumen)}", styles["Normal"]))
    elementos.append(Paragraph(f"<b>Fecha de creación:</b> {resumen.creado_en}", styles["Normal"]))
    elementos.append(Spacer(1, 16))

    for parrafo in contenido.split("\n"):
        if parrafo.strip():
            elementos.append(Paragraph(escape(parrafo.strip()), styles["BodyText"]))
            elementos.append(Spacer(1, 8))

    doc.build(elementos)

    archivo.seek(0)

    return archivo


def generar_archivo_word(resumen: Resumen) -> io.BytesIO:
    contenido = obtener_texto_resumen(resumen)

    document = Document()

    document.add_heading(resumen.titulo or "Resumen generado", level=1)

    document.add_paragraph(f"Tipo de resumen: {resumen.tipo_resumen}")
    document.add_paragraph(f"Fecha de creación: {resumen.creado_en}")

    document.add_heading("Contenido", level=2)

    for parrafo in contenido.split("\n"):
        if parrafo.strip():
            document.add_paragraph(parrafo.strip())

    archivo = io.BytesIO()
    document.save(archivo)
    archivo.seek(0)

    return archivo